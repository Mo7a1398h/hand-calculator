from docx import Document
from docx.shared import RGBColor
import os

def export_code_to_word(directory):
    # إنشاء مستند Word جديد
    doc = Document()
    doc.add_heading('Hand Calculator Project Code', 0)

    # قائمة امتدادات الملفات التي نريد تصديرها
    extensions = ['.html', '.css', '.js', '.json']

    # البحث عن جميع الملفات في المجلد
    for filename in os.listdir(directory):
        if any(filename.endswith(ext) for ext in extensions):
            file_path = os.path.join(directory, filename)
            
            # إضافة عنوان للملف
            doc.add_heading(f'File: {filename}', level=1)
            
            # قراءة محتوى الملف
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    # إضافة محتوى الملف كنص منسق
                    doc.add_paragraph(content)
                    # إضافة فاصل بين الملفات
                    doc.add_paragraph('\n' + '-'*50 + '\n')
            except Exception as e:
                doc.add_paragraph(f'Error reading file {filename}: {str(e)}')

    # حفظ المستند
    output_path = os.path.join(directory, 'hand_calculator_code.docx')
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    current_dir = os.path.dirname(os.path.abspath(__file__))
    output_file = export_code_to_word(current_dir)
    print(f'تم تصدير الكود بنجاح إلى: {output_file}')
